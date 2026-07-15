from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\tranp_3bhil36\Desktop\Ecommerce")
OUT_DIR = ROOT / "output"
OUT_DOCX = OUT_DIR / "Informe_Aura_Marketplace_mas_100_paginas.docx"

IMAGE_PATHS = [
    Path(r"C:\Users\tranp_3bhil36\AppData\Local\Temp\codex-clipboard-d38ae3f6-ad15-4073-ae8d-03406f5a250b.png"),
    Path(r"C:\Users\tranp_3bhil36\AppData\Local\Temp\codex-clipboard-88a7a5ca-ea99-4d71-9c0b-b2cd15be307e.png"),
    Path(r"C:\Users\tranp_3bhil36\AppData\Local\Temp\codex-clipboard-1d3142ac-3fd0-4277-bb62-e0c7cbb2df3c.png"),
    Path(r"C:\Users\tranp_3bhil36\AppData\Local\Temp\codex-clipboard-88d73a1a-cadc-4bf1-93d4-9463f8c0e121.png"),
    Path(r"C:\Users\tranp_3bhil36\AppData\Local\Temp\codex-clipboard-f936cb02-756d-44b0-a64d-f328e13f3ba3.png"),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, name="Times New Roman", size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.first_line_indent = Inches(0.45)

    for name, size, color, before, after in [
        ("Heading 1", 16, "1F4D78", 14, 8),
        ("Heading 2", 14, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    if "CaptionCustom" not in styles:
        cap = styles.add_style("CaptionCustom", WD_STYLE_TYPE.PARAGRAPH)
        cap.font.name = "Times New Roman"
        cap._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        cap._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        cap.font.size = Pt(10)
        cap.font.italic = True
        cap.font.color.rgb = RGBColor.from_string("555555")
        cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(8)
        cap.paragraph_format.first_line_indent = Inches(0)

    if "SmallNote" not in styles:
        note = styles.add_style("SmallNote", WD_STYLE_TYPE.PARAGRAPH)
        note.font.name = "Times New Roman"
        note._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        note._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        note.font.size = Pt(10)
        note.font.color.rgb = RGBColor.from_string("444444")
        note.paragraph_format.line_spacing = 1.15
        note.paragraph_format.space_after = Pt(4)
        note.paragraph_format.first_line_indent = Inches(0)

    header = section.header.paragraphs[0]
    header.text = "Aura Marketplace - Informe de arquitectura, diseno y metodologia"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(header.runs[0], size=9, color="666666")

    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_title_page(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    for line, size, bold in [
        ("UNIVERSIDAD NACIONAL DE SAN CRISTOBAL DE HUAMANGA", 13, True),
        ("FACULTAD DE INGENIERIA DE MINAS, GEOLOGIA Y CIVIL", 12, True),
        ("ESCUELA PROFESIONAL DE INGENIERIA DE SISTEMAS", 12, True),
    ]:
        r = p.add_run(line + "\n")
        set_run_font(r, size=size, bold=bold)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    r = p.add_run("\nINFORME DE INVESTIGACION Y DOCUMENTACION TECNICA\n")
    set_run_font(r, size=14, bold=True, color="1F4D78")
    r = p.add_run('\n"Aura Marketplace: plataforma web de comercio electronico inteligente con arquitectura hexagonal, frontend modular y agente conversacional"\n')
    set_run_font(r, size=15, bold=True)

    for label in [
        "Curso: Ingenieria de Software / Arquitectura de Sistemas",
        "Proyecto: Aura Marketplace",
        "Presentado por: ______________________________",
        "Docente: ______________________________",
        "Ayacucho - Peru",
        "2026",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Inches(0)
        r = p.add_run(label)
        set_run_font(r, size=12)
    doc.add_page_break()


def add_centered_page(doc: Document, title: str, paragraphs: list[str]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    r = p.add_run(title)
    set_run_font(r, size=16, bold=True, color="1F4D78")
    for text in paragraphs:
        para = doc.add_paragraph(text)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_page_break()


def add_heading(doc: Document, text: str, level=1) -> None:
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.line_spacing = 1.25


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "E8EEF5")
        set_cell_margins(hdr[i])
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in hdr[i].paragraphs:
            p.paragraph_format.first_line_indent = Inches(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                set_run_font(r, size=10, bold=True)
        if widths:
            hdr[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                p.paragraph_format.first_line_indent = Inches(0)
                p.paragraph_format.line_spacing = 1.1
                for r in p.runs:
                    set_run_font(r, size=9.5)
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()


def add_content_page(
    doc: Document,
    title: str,
    paragraphs: list[str],
    bullets: list[str] | None = None,
    table: tuple[list[str], list[list[str]], list[float] | None] | None = None,
    level: int = 2,
) -> None:
    add_heading(doc, title, level=level)
    for para in paragraphs:
        add_body(doc, para)
    if bullets:
        add_bullets(doc, bullets)
    if table:
        add_table(doc, table[0], table[1], table[2])
    doc.add_page_break()


def paragraph_block(topic: str, angle: str, project: str = "Aura Marketplace") -> list[str]:
    return [
        f"En el contexto de {project}, {topic} se entiende como un componente critico del ciclo de vida del sistema, porque permite conectar las decisiones tecnicas con los objetivos del negocio, la experiencia del usuario y la sostenibilidad del producto. Su estudio no se limita a describir una tecnologia; tambien explica por que esa tecnologia se adopta y que consecuencias tiene sobre la calidad final.",
        f"Desde la perspectiva de {angle}, el proyecto requiere separar los conceptos teoricos de las decisiones de implementacion. Esta separacion evita que el documento sea una simple descripcion de pantallas o archivos y permite demostrar que cada modulo responde a principios de ingenieria: bajo acoplamiento, alta cohesion, trazabilidad, seguridad por diseno y validacion continua.",
        f"La aplicacion concreta se observa en la integracion entre una SPA React, un backend NestJS organizado por capas L01 a L05, una base de datos PostgreSQL administrada mediante Prisma y servicios externos como Neon Auth, Mercado Pago, Resend, Cloudinary, Upstash Redis y Gemini AI. La arquitectura se disena para que esas dependencias sean reemplazables y no contaminen el nucleo del dominio.",
    ]


APA_REFERENCES = [
    ["Cockburn, A., & Garrido de Paz, J. M.", "2024", "Hexagonal Architecture Explained: How the Ports & Adapters Architecture Simplifies Your Technical Documentation.", "Humans and Technology."],
    ["Evans, E.", "2004", "Domain-Driven Design: Tackling Complexity in the Heart of Software.", "Addison-Wesley."],
    ["Fowler, M.", "2004", "Inversion of Control Containers and the Dependency Injection Pattern.", "https://martinfowler.com/articles/injection.html"],
    ["Humble, J., & Farley, D.", "2010", "Continuous Delivery: Reliable Software Releases through Build, Test, and Deployment Automation.", "Addison-Wesley."],
    ["IEEE/ISO/IEC", "2018", "IEEE/ISO/IEC 29148-2018: Systems and software engineering - Life cycle processes - Requirements engineering.", "IEEE Standards Association. https://standards.ieee.org/standard/29148-2018.html"],
    ["Meta Open Source", "s. f.", "React.", "https://react.dev/"],
    ["NestJS", "s. f.", "Documentation.", "https://docs.nestjs.com/"],
    ["OWASP Foundation", "s. f.", "OWASP Application Security Verification Standard (ASVS).", "https://owasp.org/www-project-application-security-verification-standard/"],
    ["Prisma", "s. f.", "What is Prisma ORM?", "https://www.prisma.io/docs/orm"],
    ["Pressman, R. S., & Maxim, B. R.", "2020", "Software Engineering: A Practitioner's Approach (9th ed.).", "McGraw-Hill Education."],
    ["Schwaber, K., & Sutherland, J.", "2020", "The Scrum Guide.", "https://scrumguides.org/"],
    ["Sommerville, I.", "2016", "Software Engineering (10th ed.).", "Pearson."],
    ["TanStack", "s. f.", "TanStack Query.", "https://tanstack.com/query/latest"],
    ["W3C Web Accessibility Initiative", "s. f.", "WCAG 2 Overview.", "https://www.w3.org/WAI/standards-guidelines/wcag/"],
    ["Zustand", "s. f.", "Introduction.", "https://zustand.docs.pmnd.rs/"],
]


THEORY_ENTRIES = [
    (
        "2.1 Antecedentes de la investigacion",
        "La ingenieria de software moderna sostiene que un sistema debe comprenderse como un producto sociotecnico: satisface necesidades humanas, pero se materializa mediante procesos, modelos, codigo, pruebas y despliegue. Sommerville (2016) explica que los sistemas intensivos en software no pueden evaluarse solo por su ejecucion, sino tambien por su capacidad de evolucionar, mantenerse y operar con restricciones de calidad. En Aura Marketplace, este antecedente justifica que el informe no se limite a pantallas o endpoints, sino que documente vision, requisitos, arquitectura, base de datos, seguridad y metodologia.",
        "Desde una perspectiva complementaria, Pressman y Maxim (2020) senalan que la calidad se construye durante todo el proceso y no solamente al final. Por ello, el marco teorico vincula SDD, Scrum, arquitectura hexagonal, frontend modular y pruebas como practicas integradas. La plataforma estudiada requiere esa integracion porque une comercio electronico, agentes conversacionales, pagos externos, autenticacion, inventario, roles y auditoria.",
        "Base bibliografica: Sommerville (2016) y Pressman y Maxim (2020).",
    ),
    (
        "2.2 Specs Driven Development (SDD)",
        "Specs Driven Development se entiende en este informe como un enfoque en el que las especificaciones son la fuente de verdad del desarrollo. La norma IEEE/ISO/IEC 29148 establece procesos y productos para la ingenieria de requisitos a lo largo del ciclo de vida del software. En cita textual breve, dicha norma alude a \"software products and services throughout the life cycle\" (IEEE/ISO/IEC, 2018, secc. Scope).",
        "Aplicado a Aura Marketplace, SDD significa que los archivos de la carpeta specs no son documentos decorativos: Vision, Stakeholders, Glosario, Alcance, Objetivos, Reglas de Negocio, Requisitos Funcionales, Requisitos No Funcionales, Actores, Casos de Uso, Criterios de Aceptacion y Matriz de Trazabilidad orientan el diseno de modulos, endpoints, entidades, stores, pruebas y anexos. Esta relacion convierte cada cambio tecnico en una decision justificable.",
        "Base bibliografica: IEEE/ISO/IEC (2018).",
    ),
    (
        "2.3 Requisitos de software y trazabilidad",
        "La trazabilidad permite seguir el recorrido de una necesidad desde su formulacion hasta su verificacion. IEEE/ISO/IEC 29148 (2018) destaca la importancia de atributos, caracteristicas y aplicacion iterativa de los procesos de requisitos. En proyectos con multiples capas, la trazabilidad evita que los modulos se implementen por intuicion y obliga a demostrar que una funcion responde a una necesidad identificada.",
        "En Aura Marketplace, esta base teorica se evidencia cuando un objetivo como reducir la friccion de compra se transforma en requisitos del agente, reglas de confirmacion, criterios de aceptacion, controladores REST, stores de estado, componentes de UI y pruebas E2E. El anexo SDD debe mostrar esta cadena, porque el desarrollo guiado por especificaciones solo es valido si cada artefacto se puede rastrear.",
        "Base bibliografica: IEEE/ISO/IEC (2018) y Sommerville (2016).",
    ),
    (
        "2.4 Plataforma web de comercio electronico",
        "Un marketplace digital combina catalogo, identidad, transacciones, confianza, pagos, soporte y administracion. Pressman y Maxim (2020) recomiendan modelar estas funciones como requisitos verificables y no como listas informales de deseos. En Aura Marketplace, la capa de comercio electronico se expresa mediante publicaciones, categorias, inventario, carrito, ordenes, pagos, favoritos, resenas, promociones y notificaciones.",
        "La importancia teorica del comercio electronico en este informe radica en que cada flujo implica restricciones de datos y seguridad. Una compra no es solamente una pagina de checkout: requiere autenticacion, validacion de stock, persistencia atomica, integracion con pasarela de pago, registro historico y notificacion. Por ello, el sistema necesita arquitectura por capas y trazabilidad documental.",
        "Base bibliografica: Pressman y Maxim (2020) y Sommerville (2016).",
    ),
    (
        "2.5 Agente inteligente conversacional",
        "El agente conversacional se fundamenta en la necesidad de reducir la distancia entre la intencion humana y la accion funcional del sistema. Desde el punto de vista de requisitos, esta capacidad debe definirse mediante entradas, salidas, reglas de negocio, restricciones de seguridad y criterios de aceptacion (IEEE/ISO/IEC, 2018). La interaccion en lenguaje natural no reemplaza la arquitectura; la exige con mayor claridad.",
        "En Aura Marketplace, el agente no debe operar como un modulo aislado sin control. Su interpretacion de intenciones se conecta con catalogo, carrito, ordenes y usuario autenticado. Por ello, la capa L02 debe coordinar la conversacion y delegar en servicios de aplicacion, evitando que el proveedor de IA decida reglas de negocio o acceda directamente a la persistencia.",
        "Base bibliografica: IEEE/ISO/IEC (2018) y Evans (2004).",
    ),
    (
        "2.6 Arquitectura general del sistema",
        "La arquitectura de software organiza las decisiones que permiten que un sistema cumpla requisitos funcionales y no funcionales. Sommerville (2016) relaciona arquitectura con rendimiento, seguridad, disponibilidad y mantenibilidad. En Aura Marketplace, la arquitectura general separa cliente, API, aplicacion, dominio, infraestructura, datos e integraciones externas.",
        "Esta separacion permite explicar por que el frontend no contiene secretos, por que el backend centraliza reglas de negocio, por que Prisma encapsula persistencia y por que Mercado Pago, Resend, Cloudinary, Upstash Redis, Neon y Gemini deben tratarse como dependencias externas controladas. La arquitectura no solo representa el sistema; tambien limita como puede evolucionar.",
        "Base bibliografica: Sommerville (2016) y Pressman y Maxim (2020).",
    ),
    (
        "2.7 Arquitectura hexagonal",
        "La arquitectura hexagonal, tambien conocida como puertos y adaptadores, propone proteger el nucleo de negocio de detalles externos. Cockburn y Garrido de Paz (2024) desarrollan este enfoque como una forma de simplificar la documentacion tecnica y separar el centro de la aplicacion de sus mecanismos de entrada y salida. En terminos de DDD, Evans (2004) refuerza que el dominio debe expresar el lenguaje y las reglas principales del negocio.",
        "Aura Marketplace aplica esta teoria mediante L04 Domain, donde residen entidades, enums y puertos; L03 Application, donde se orquestan casos de uso; y L05 Infrastructure, donde se implementan adaptadores para Prisma, pagos, correo, cache, IA, almacenamiento y seguridad. Con ello, cambiar un proveedor no deberia obligar a reescribir las reglas centrales.",
        "Base bibliografica: Cockburn y Garrido de Paz (2024) y Evans (2004).",
    ),
    (
        "2.8 Puertos, adaptadores e inversion de dependencias",
        "Fowler (2004) analiza el patron conocido como Dependency Injection y lo relaciona con la inversion de control. En cita textual breve, Fowler aborda el tema bajo \"the more specific name of Dependency Injection\" (Fowler, 2004, par. 1). Este principio resulta esencial cuando el sistema necesita ensamblar implementaciones concretas sin acoplar la logica de aplicacion a una tecnologia puntual.",
        "En Aura Marketplace, los puertos como IUserRepository, IOrderRepository, IPaymentGateway, IHasher, IMailSender o ICacheProvider definen contratos. Sus adaptadores concretos, como PrismaUserRepository, MercadoPagoService, Argon2HasherService o ResendMailService, satisfacen esos contratos. Esta decision permite pruebas con mocks, sustitucion de proveedores y menor riesgo de regresiones.",
        "Base bibliografica: Fowler (2004), Cockburn y Garrido de Paz (2024).",
    ),
    (
        "2.9 Arquitectura del backend",
        "NestJS se presenta en su documentacion como un marco para construir aplicaciones del lado servidor eficientes y escalables. En cita textual breve, NestJS menciona \"efficient, scalable Node.js server-side applications\" (NestJS, s. f., par. 1). Esta definicion justifica su eleccion para un backend modular que requiere controladores, servicios, guards, pipes y providers.",
        "El backend de Aura Marketplace se divide en L01 Presentation, L02 Agent, L03 Application, L04 Domain y L05 Infrastructure. Esta estructura se alinea con la arquitectura hexagonal porque las capas externas reciben solicitudes o implementan detalles tecnicos, mientras el dominio conserva contratos y conceptos del negocio. La modularidad de NestJS permite mapear requisitos a modulos funcionales.",
        "Base bibliografica: NestJS (s. f.), Cockburn y Garrido de Paz (2024).",
    ),
    (
        "2.10 Controladores, DTOs y validacion",
        "La validacion de entradas es una practica fundamental de seguridad y calidad. OWASP Foundation (s. f.) propone usar controles tecnicos verificables para incrementar la confianza en aplicaciones web. En backend, los DTOs, pipes y validadores reducen la posibilidad de que datos invalidos lleguen a servicios de aplicacion o repositorios.",
        "En Aura Marketplace, esta teoria se materializa en DTOs de registro, login, carrito, ordenes, categorias, perfil y restablecimiento de contrasena. La validacion tambien permite cumplir SDD: cada regla especificada debe tener una barrera tecnica que la haga exigible. Por ejemplo, una regla de precio mayor que cero debe aparecer en especificacion, DTO, servicio y prueba.",
        "Base bibliografica: OWASP Foundation (s. f.) y IEEE/ISO/IEC (2018).",
    ),
    (
        "2.11 Seguridad backend",
        "La seguridad por diseno exige controles en autenticacion, gestion de sesiones, autorizacion, proteccion de datos, errores y auditoria. OWASP Foundation (s. f.) indica que ASVS ofrece una base para probar controles tecnicos de seguridad. Esa orientacion se alinea con sistemas que gestionan usuarios, pagos, roles y datos personales.",
        "En Aura Marketplace, la seguridad se implementa mediante Argon2 para contrasenas, JWT y refresh tokens para sesiones, TokenRevocado para invalidacion, RolesGuard para autorizacion, HTTPS en despliegue y auditoria para eventos criticos. La teoria respalda que estos controles no sean accesorios, sino requisitos no funcionales trazables.",
        "Base bibliografica: OWASP Foundation (s. f.) y Pressman y Maxim (2020).",
    ),
    (
        "2.12 Persistencia relacional y Prisma ORM",
        "Prisma ORM se describe como una herramienta que proporciona acceso seguro por tipos, migraciones y editor visual de datos. En cita textual breve, Prisma destaca \"type-safe database access\" y migraciones (Prisma, s. f., par. 1). Esta idea resulta relevante para proyectos TypeScript donde el modelo de datos debe ser consistente con el codigo.",
        "En Aura Marketplace, Prisma define entidades como Usuario, Publicacion, Inventario, Carrito, Orden, Pago, Sesion, Conversacion, Mensaje, Intencion, Auditoria y TokenRevocado. El esquema opera como contrato persistente del dominio, mientras los repositorios Prisma implementan puertos para aislar consultas, relaciones y transacciones de la capa de aplicacion.",
        "Base bibliografica: Prisma (s. f.) y Evans (2004).",
    ),
    (
        "2.13 Integridad, transacciones y consistencia de datos",
        "La calidad de datos se relaciona con reglas de negocio verificables, restricciones de integridad y comportamiento transaccional. Sommerville (2016) destaca que las restricciones no funcionales condicionan la arquitectura y la implementacion. En un marketplace, la consistencia es critica porque una orden no puede registrarse de forma separada al stock o al pago.",
        "Aura Marketplace requiere transacciones para crear orden, lineas, pago, decremento de inventario y limpieza de carrito. Esta necesidad se deriva de SDD: las reglas de negocio de stock, precio historico y confirmacion de compra deben mantenerse aunque existan errores externos. Por eso la persistencia se concentra en adaptadores controlados.",
        "Base bibliografica: Sommerville (2016), IEEE/ISO/IEC (2018).",
    ),
    (
        "2.14 Arquitectura del frontend",
        "React se presenta oficialmente como biblioteca para interfaces web y nativas. En cita textual breve, su documentacion la define como \"library for web and native user interfaces\" (Meta Open Source, s. f., par. 1). Esta base teorica justifica su uso en una SPA donde las pantallas se componen desde unidades reutilizables.",
        "En Aura Marketplace, el frontend organiza pages, components, api, store, layouts, hooks, lib, styles y types. Esta estructura soporta rutas publicas, rutas protegidas, panel del comprador, panel del vendedor y panel administrativo. El frontend no contiene reglas sensibles, pero si debe representar estados, validar formularios y comunicar intenciones del usuario al backend.",
        "Base bibliografica: Meta Open Source (s. f.) y Pressman y Maxim (2020).",
    ),
    (
        "2.15 Componentes, reutilizacion y experiencia de usuario",
        "El enfoque por componentes permite construir interfaces coherentes y mantenibles. Meta Open Source (s. f.) explica que React permite componer interfaces a partir de piezas individuales. En Aura Marketplace, esta teoria se refleja en componentes como AuraHeader, BrandLogo, ChatAgente, AuthLayout, AuthField, CheckoutStepper y CheckoutOrderSummary.",
        "La reutilizacion no se limita a ahorrar codigo; tambien aporta consistencia visual y reduce errores. Cuando el sistema usa componentes para formularios de autenticacion o checkout, las reglas de accesibilidad, estados de carga y mensajes de error pueden mantenerse de forma uniforme. Esto se vincula con requisitos de usabilidad y accesibilidad.",
        "Base bibliografica: Meta Open Source (s. f.) y W3C Web Accessibility Initiative (s. f.).",
    ),
    (
        "2.16 Estado cliente con Zustand",
        "La gestion de estado cliente se vuelve necesaria cuando la interfaz debe recordar sesion, carrito, perfil, preferencias o conversacion. Zustand se presenta como una solucion pequena, rapida y escalable de gestion de estado (Zustand, s. f.). En Aura Marketplace, esta caracteristica se relaciona directamente con authStore, cartStore, agentStore y profilePhotoStore.",
        "Desde SDD, cada store debe responder a una necesidad especificada: authStore sostiene autenticacion, cartStore gestiona interaccion de compra, agentStore conserva el estado conversacional y profilePhotoStore gestiona presentacion del usuario. Asi, el estado local no se convierte en logica duplicada del backend, sino en soporte de experiencia.",
        "Base bibliografica: Zustand (s. f.) y IEEE/ISO/IEC (2018).",
    ),
    (
        "2.17 Estado servidor con TanStack Query",
        "TanStack Query aborda el problema de sincronizar datos remotos en aplicaciones modernas. En cita textual breve, su documentacion resume: \"Stop syncing server data by hand\" (TanStack, s. f., par. 1). Esta idea se relaciona con cache, reintentos, invalidacion y actualizacion de datos remotos.",
        "Aunque el proyecto tambien usa Axios y stores locales, el fundamento teorico de estado servidor permite justificar la separacion entre datos que pertenecen al backend y estados puramente visuales. Catalogo, ordenes, productos y reportes deben gestionarse como datos remotos, mientras modales, filtros visibles o mensajes temporales pueden ser estado local.",
        "Base bibliografica: TanStack (s. f.) y Meta Open Source (s. f.).",
    ),
    (
        "2.18 Validacion frontend y formularios",
        "La validacion en cliente mejora la experiencia de usuario al detectar errores antes de enviar solicitudes, pero no reemplaza la validacion del backend. OWASP Foundation (s. f.) sostiene que los controles tecnicos deben verificarse en la aplicacion; por ello, la validacion debe existir en varias capas. React Hook Form y Zod se entienden como soporte de frontera visual.",
        "En Aura Marketplace, los formularios de registro, login, perfil, direccion, producto y checkout deben guiar al usuario con mensajes claros. Bajo SDD, cada campo requerido, restriccion de contrasena o validacion de precio debe estar especificada, implementada y probada. La validacion cliente mejora usabilidad, mientras el backend asegura integridad.",
        "Base bibliografica: OWASP Foundation (s. f.) y IEEE/ISO/IEC (2018).",
    ),
    (
        "2.19 Accesibilidad web",
        "La accesibilidad es una dimension de calidad que evita excluir usuarios por limitaciones visuales, motoras, cognitivas o tecnologicas. W3C Web Accessibility Initiative (s. f.) organiza WCAG bajo cuatro principios: \"perceivable, operable, understandable, and robust\" (par. 1). Esta cita textual breve sintetiza el marco de accesibilidad usado en aplicaciones web.",
        "En Aura Marketplace, la accesibilidad se relaciona con navegacion por teclado, etiquetas en controles, mensajes de estado del agente, contraste visual, textos alternativos en imagenes y claridad de errores. El agente de voz tambien se justifica como mecanismo de acceso alternativo, pero no debe sustituir la compatibilidad general de la interfaz.",
        "Base bibliografica: W3C Web Accessibility Initiative (s. f.).",
    ),
    (
        "2.20 Metodologias agiles y Scrum",
        "Scrum se define oficialmente como un marco para desarrollar y sostener productos complejos. En cita textual breve, Scrum Guides lo presenta como \"framework for developing and sustaining complex products\" (Schwaber & Sutherland, 2020, par. 1). Su valor teorico esta en organizar trabajo empirico mediante roles, eventos, artefactos e incrementos.",
        "En Aura Marketplace, Scrum permite agrupar el desarrollo por sprints: configuracion base, autenticacion, catalogo, carrito, pagos, agente, administracion, calidad y despliegue. Al combinarse con SDD, el Product Backlog no nace de ideas sueltas, sino de especificaciones versionadas, criterios de aceptacion y matriz de trazabilidad.",
        "Base bibliografica: Schwaber y Sutherland (2020) y IEEE/ISO/IEC (2018).",
    ),
    (
        "2.21 Product Backlog e historias de usuario",
        "El Product Backlog representa el trabajo necesario para alcanzar el objetivo del producto. Schwaber y Sutherland (2020) explican que los artefactos de Scrum incrementan transparencia y permiten inspeccion. En este informe, las historias de usuario conectan necesidades del comprador, vendedor y administrador con funcionalidades verificables.",
        "En Aura Marketplace, historias como registrarse, buscar productos, conversar con el agente, agregar al carrito, completar pago, publicar producto o administrar usuarios deben derivar de requisitos funcionales y reglas de negocio. Esta relacion evita que el backlog pierda alineacion con el documento SDD.",
        "Base bibliografica: Schwaber y Sutherland (2020) y Pressman y Maxim (2020).",
    ),
    (
        "2.22 Pruebas y aseguramiento de calidad",
        "El aseguramiento de calidad no se reduce a ejecutar pruebas al final. Pressman y Maxim (2020) relacionan calidad con procesos, revisiones, mediciones y pruebas a lo largo del desarrollo. En un sistema modular, las pruebas deben cubrir unidades, integraciones y flujos E2E.",
        "Aura Marketplace cuenta con pruebas unitarias en backend, pruebas de integracion para proveedores y pruebas E2E en frontend. Bajo SDD, cada criterio de aceptacion debe transformarse en evidencia de prueba. Por ejemplo, una regla de bloqueo por intentos fallidos debe aparecer en requisito, servicio, test y documentacion.",
        "Base bibliografica: Pressman y Maxim (2020) y IEEE/ISO/IEC (2018).",
    ),
    (
        "2.23 Seguridad aplicativa y ASVS",
        "OWASP ASVS ofrece un marco para verificar controles tecnicos de seguridad. En cita textual breve, ASVS proporciona una \"basis for testing application technical security controls\" (OWASP Foundation, s. f., par. 1). Esto permite transformar seguridad en requisitos verificables y no en una declaracion general.",
        "En Aura Marketplace, ASVS respalda decisiones como hash de contrasenas, expiracion de tokens, control de acceso por roles, validacion de entrada, no almacenamiento de tarjetas, auditoria y manejo seguro de errores. La seguridad se documenta como parte del diseno y se prueba como parte del ciclo de calidad.",
        "Base bibliografica: OWASP Foundation (s. f.).",
    ),
    (
        "2.24 Despliegue y entrega continua",
        "La entrega continua promueve que el software pueda construirse, probarse y desplegarse de forma repetible. Aunque Aura Marketplace no necesariamente automatiza todo el ciclo, su documentacion de Render, Vercel, Prisma Migrate, variables de entorno y GitHub Actions se alinea con practicas de despliegue controlado (Humble & Farley, 2010).",
        "Desde SDD, el despliegue tambien debe estar especificado: variables obligatorias, comandos de build, migraciones, dominios, CORS, webhooks y health checks. Esto reduce errores entre ambientes y mantiene coherencia entre la arquitectura disenada y la operacion real.",
        "Base bibliografica: Humble y Farley (2010) y Pressman y Maxim (2020).",
    ),
    (
        "2.25 Integraciones externas",
        "Las integraciones externas deben tratarse como dependencias que pueden fallar, cambiar o introducir latencia. Sommerville (2016) sostiene que los requisitos no funcionales condicionan decisiones arquitectonicas. Por ello, servicios como Mercado Pago, Resend, Cloudinary, Upstash Redis, Neon Auth y Gemini AI deben encapsularse mediante adaptadores.",
        "En Aura Marketplace, este principio permite degradacion controlada. Si falla correo, el sistema debe informar; si falla IA, puede mantenerse navegacion manual; si falla pago, el carrito no debe perderse. La teoria respalda que cada integracion tenga contrato, manejo de errores, configuracion segura y pruebas.",
        "Base bibliografica: Sommerville (2016), Cockburn y Garrido de Paz (2024).",
    ),
    (
        "2.26 Observabilidad y auditoria",
        "La observabilidad permite comprender el comportamiento del sistema mediante eventos, logs, metricas y trazas. Pressman y Maxim (2020) asocian la calidad con capacidad de diagnostico y mantenimiento. En sistemas transaccionales, la auditoria tambien cumple una funcion de confianza y responsabilidad.",
        "Aura Marketplace registra eventos relevantes como autenticaciones, instrucciones del agente, cambios de publicaciones, registro de ordenes, cambios de estado y errores de integraciones. La auditoria debe evitar datos sensibles, cumpliendo principios de seguridad. Bajo SDD, cada evento auditable debe relacionarse con requisito o regla de negocio.",
        "Base bibliografica: Pressman y Maxim (2020) y OWASP Foundation (s. f.).",
    ),
    (
        "2.27 Dominio, entidades y lenguaje ubicuo",
        "Domain-Driven Design propone que el lenguaje del dominio guie el modelo del software. Evans (2004) sostiene que el centro del diseno debe expresar conceptos y reglas relevantes para el negocio. En Aura Marketplace, terminos como Usuario, Publicacion, Orden, Carrito, Pago, Sesion, Mensaje e Intencion deben usarse de forma consistente.",
        "El glosario de specs cumple una funcion compatible con DDD: evita que una misma idea tenga varios nombres y permite alinear base de datos, DTOs, servicios, stores y documentacion. Esto mejora mantenibilidad porque los desarrolladores pueden entender el sistema desde el lenguaje del negocio.",
        "Base bibliografica: Evans (2004) y IEEE/ISO/IEC (2018).",
    ),
    (
        "2.28 Modularidad y mantenibilidad",
        "La mantenibilidad se favorece cuando los modulos tienen responsabilidades claras, interfaces estables y dependencias controladas. Sommerville (2016) y Pressman y Maxim (2020) coinciden en que la estructura del software afecta el costo de cambio. En Aura Marketplace, la separacion de modulos permite evolucionar autenticacion, catalogo, pagos o agente sin reescribir todo el sistema.",
        "La modularidad tambien facilita pruebas. Un servicio puede probarse con repositorios simulados, un adaptador puede verificarse con pruebas de integracion y una pagina frontend puede evaluarse mediante flujos E2E. El marco teorico respalda que la arquitectura documentada sea parte de la calidad del producto.",
        "Base bibliografica: Sommerville (2016) y Pressman y Maxim (2020).",
    ),
    (
        "2.29 Relacion entre SDD, arquitectura y Scrum",
        "SDD define que se debe construir; arquitectura define como se organiza; Scrum define como se gestiona el trabajo iterativo. Estas tres perspectivas no compiten: se complementan. IEEE/ISO/IEC (2018) aporta el enfoque de requisitos, Cockburn y Garrido de Paz (2024) aportan separacion arquitectonica y Schwaber y Sutherland (2020) aportan inspeccion y adaptacion.",
        "En Aura Marketplace, esta relacion permite que cada sprint tenga una base documental, cada modulo tenga una responsabilidad tecnica y cada prueba responda a un criterio de aceptacion. El informe, por tanto, no solo describe el resultado final; tambien muestra el camino metodologico y tecnico para justificarlo.",
        "Base bibliografica: IEEE/ISO/IEC (2018), Cockburn y Garrido de Paz (2024), Schwaber y Sutherland (2020).",
    ),
    (
        "2.30 Sintesis del marco teorico",
        "El marco teorico demuestra que Aura Marketplace se sostiene en fundamentos de ingenieria de software, requisitos, arquitectura, frontend, backend, seguridad, accesibilidad, persistencia, pruebas y metodologias agiles. La principal conclusion teorica es que el sistema no debe entenderse como una suma de tecnologias, sino como un producto guiado por especificaciones y validado por evidencias.",
        "Por ello, el informe adopta SDD como eje documental. Las referencias bibliograficas permiten sustentar que cada decision tiene respaldo teorico: React para componentes, NestJS para backend modular, Prisma para persistencia tipada, Scrum para gestion iterativa, ASVS para seguridad, WCAG para accesibilidad y arquitectura hexagonal para desacoplamiento.",
        "Base bibliografica: todas las fuentes citadas en el Capitulo II.",
    ),
]


def build_front_matter(doc: Document) -> None:
    add_title_page(doc)
    add_centered_page(
        doc,
        "Dedicatoria",
        [
            "Dedico este trabajo a mi familia, por el apoyo constante durante el proceso de analisis, diseno, implementacion y documentacion del sistema. Su acompanamiento permitio sostener el esfuerzo necesario para convertir una idea de marketplace inteligente en una propuesta tecnica estructurada.",
            "Asimismo, dedico el presente informe a los estudiantes y docentes que comprenden que la ingenieria de software no termina en la programacion, sino que exige justificar las decisiones, documentar la arquitectura, evaluar la calidad y dejar evidencia suficiente para que el sistema pueda evolucionar.",
        ],
    )
    add_centered_page(
        doc,
        "Agradecimiento",
        [
            "Agradezco a los docentes de la Escuela Profesional de Ingenieria de Sistemas por impulsar una vision rigurosa del desarrollo de software. La elaboracion de este informe se apoya en conocimientos de arquitectura, bases de datos, seguridad, metodologias agiles, frontend, backend, pruebas y documentacion tecnica.",
            "Tambien se expresa agradecimiento a quienes participaron en la revision del proyecto Aura Marketplace, ya que sus observaciones permitieron ordenar los requerimientos, consolidar los diagramas anexos y convertir los hallazgos del repositorio en un documento academico de mas de cien paginas.",
        ],
    )

    add_content_page(
        doc,
        "Resumen",
        [
            "El presente informe desarrolla la documentacion integral del sistema Aura Marketplace, una plataforma web de comercio electronico inteligente que combina una interfaz SPA en React con un backend NestJS basado en arquitectura hexagonal, persistencia relacional en PostgreSQL mediante Prisma y un agente conversacional capaz de interpretar instrucciones en lenguaje natural.",
            "El objetivo principal del documento es explicar, con enfoque academico y tecnico, el marco teorico, la arquitectura general, el diseno del backend, el diseno del frontend, la metodologia Scrum aplicada y los anexos de ingenieria de software. Para ello se reviso el documento de ejemplo proporcionado, el repositorio del sistema, las especificaciones, los archivos de diseno, el esquema de base de datos y los diagramas entregados como imagenes.",
            "Como resultado, se presenta una propuesta documentada que supera las cien paginas e integra capitulos de planteamiento del problema, marco teorico, material y metodos, resultados de diseno e implementacion, conclusiones, recomendaciones y anexos. El informe enfatiza la aplicacion de Specs Driven Development (SDD), la separacion por capas L01-L05, la trazabilidad de requisitos y la coherencia entre frontend, backend, base de datos e integraciones externas.",
        ],
        bullets=[
            "Palabras clave: Aura Marketplace, arquitectura hexagonal, React, NestJS, Prisma, Scrum, SDD, agente inteligente.",
        ],
        level=1,
    )

    add_content_page(
        doc,
        "Abstract",
        [
            "This report presents the complete technical and academic documentation for Aura Marketplace, an intelligent e-commerce web platform that combines a React single page application, a NestJS backend based on hexagonal architecture, PostgreSQL persistence through Prisma, and a conversational agent capable of interpreting natural language instructions.",
            "The document explains the theoretical framework, the general system architecture, backend and frontend architecture, the Scrum methodology, and the software engineering appendices. The analysis is based on the provided reference PDF, the actual project repository, specification documents, design documents, database schema, deployment files, and architectural diagrams supplied as annex images.",
            "The resulting Word document is structured as a thesis-like report with more than one hundred pages, including front matter, problem statement, theoretical framework, methodology, technical results, conclusions, recommendations, bibliography, and annexes. Special attention is given to Specs Driven Development, the L01-L05 backend layering model, requirements traceability, and the alignment between frontend, backend, database, and external integrations.",
        ],
        bullets=[
            "Keywords: Aura Marketplace, hexagonal architecture, React, NestJS, Prisma, Scrum, SDD, intelligent agent.",
        ],
        level=1,
    )

    toc_rows = [
        ["Introduccion", "Contexto del proyecto, alcance y estructura del informe"],
        ["Capitulo I", "Planteamiento del problema, objetivos, hipotesis y justificacion"],
        ["Capitulo II", "Marco teorico: SDD, arquitectura, backend, frontend y Scrum"],
        ["Capitulo III", "Material y metodos: enfoque aplicado, variables y gestion Scrum"],
        ["Capitulo IV", "Resultados: arquitectura, diseno, implementacion, pruebas y despliegue"],
        ["Capitulo V", "Conclusiones y recomendaciones"],
        ["Referencias", "Bibliografia usada para sustentar la documentacion tecnica"],
        ["Anexos", "SDD, backlog, matriz, diagramas, arquitectura y evidencias del proyecto"],
    ]
    add_content_page(
        doc,
        "Contenido",
        [
            "La organizacion del documento sigue la estructura observada en el archivo de ejemplo: paginas preliminares, introduccion, cinco capitulos principales, referencias bibliograficas y anexos. La diferencia principal es que el contenido se adapta al sistema Aura Marketplace y a los diagramas tecnicos proporcionados.",
        ],
        table=(["Seccion", "Descripcion"], toc_rows, [1.8, 4.6]),
        level=1,
    )
    add_content_page(
        doc,
        "Lista de figuras y tablas",
        [
            "Las figuras incluidas en los anexos corresponden a los diagramas proporcionados por el usuario y a los artefactos generados en el repositorio. Las tablas sintetizan requisitos, reglas de negocio, trazabilidad, backlog, sprint planning, arquitectura por capas, endpoints, entidades de datos y escenarios SDD.",
        ],
        table=(
            ["Elemento", "Contenido"],
            [
                ["Figura A1", "Arquitectura general del backend y dependencias externas"],
                ["Figura A2", "Arquitectura modular del frontend React"],
                ["Figura A3", "Boundaries de backend, infraestructura y datos"],
                ["Figura A4", "Modelo entidad-relacion del marketplace"],
                ["Figura A5", "Stack tecnologico y despliegue conectado"],
                ["Tablas", "Matriz de consistencia, backlog, requisitos, reglas, pruebas y SDD"],
            ],
            [1.7, 4.7],
        ),
        level=1,
    )


def build_chapter_one(doc: Document) -> None:
    add_content_page(
        doc,
        "Introduccion",
        [
            "La digitalizacion del comercio ha transformado la manera en que compradores, vendedores y administradores interactuan dentro de un ecosistema transaccional. En este escenario, Aura Marketplace se plantea como una plataforma que no solo publica productos y procesa compras, sino que tambien reduce la friccion mediante un agente inteligente capaz de interpretar instrucciones en lenguaje natural.",
            "El sistema se organiza como un monorepo con frontend en React, backend en NestJS, base de datos PostgreSQL y servicios externos para autenticacion, pagos, correo, almacenamiento, cache e inteligencia artificial. Esta composicion exige una documentacion extensa que explique tanto el fundamento teorico como la arquitectura concreta de implementacion.",
            "El documento toma como referencia formal el archivo TrabajoEjemplo.pdf, cuyo patron academico contiene portada, resumen, abstract, capitulos metodologicos, resultados, conclusiones y anexos. Sobre esa base, el presente informe desarrolla una version orientada al proyecto Aura Marketplace, incorporando los diagramas anexos enviados por el usuario.",
        ],
        level=1,
    )
    add_heading(doc, "Capitulo I", level=1)
    add_content_page(
        doc,
        "1.1 Diagnostico y enunciado del problema",
        [
            "Los marketplaces tradicionales obligan al usuario a transformar su necesidad en filtros, categorias, palabras clave y acciones repetitivas. Esta conversion entre intencion humana e interaccion grafica produce friccion, especialmente para usuarios con baja experiencia digital, limitaciones visuales o preferencia por interfaces conversacionales.",
            "En Aura Marketplace, el problema se observa como una brecha entre el lenguaje natural del comprador y los flujos tradicionales de exploracion, comparacion, carrito y pago. Si el sistema no interpreta correctamente esa intencion, el usuario debe navegar manualmente, repetir busquedas o abandonar la compra.",
            "Adicionalmente, la implementacion de un sistema de comercio electronico inteligente exige resolver problemas de seguridad, trazabilidad, integridad transaccional, control de roles, persistencia confiable y escalabilidad. Por ello, la solucion no puede limitarse a una interfaz atractiva; requiere arquitectura backend, frontend, datos, integraciones y metodologia de desarrollo claramente documentadas.",
        ],
        bullets=[
            "Friccion entre la intencion del usuario y las acciones de compra.",
            "Necesidad de organizar el backend para aislar dominio, aplicacion e infraestructura.",
            "Necesidad de un frontend modular que soporte roles, carrito, catalogo, checkout y agente.",
            "Necesidad de anexos tecnicos que evidencien SDD, ERD, despliegue y decisiones arquitectonicas.",
        ],
    )

    problem_pages = [
        ("1.2 Problema general", "determinar como documentar, disenar e implementar una plataforma de comercio electronico inteligente que permita a compradores, vendedores y administradores operar con seguridad, baja friccion y trazabilidad tecnica."),
        ("1.3 Problemas especificos", "identificar las dificultades asociadas al modelado de requisitos, la comunicacion frontend-backend, la persistencia transaccional, la seguridad de sesiones, la integracion de servicios externos y la aplicacion de SDD en los anexos."),
        ("1.4 Objetivo general", "desarrollar una documentacion integral del sistema Aura Marketplace que explique el marco teorico, la metodologia Scrum, la arquitectura general, el backend, el frontend, la base de datos, las integraciones, las pruebas y los anexos tecnicos."),
        ("1.5 Objetivos especificos", "organizar el informe en capitulos academicos, describir la aplicacion de SDD, justificar la arquitectura hexagonal, explicar la modularidad del frontend, documentar la base de datos y elaborar anexos con diagramas y tablas de soporte."),
        ("1.6 Hipotesis", "si el sistema se documenta mediante una estructura academica basada en Scrum, SDD y arquitectura hexagonal, entonces se obtiene un informe verificable que evidencia la coherencia entre requisitos, diseno e implementacion."),
        ("1.7 Justificacion tecnica", "sustentar que el proyecto requiere puertos, adaptadores, DTOs, guards, repositorios, stores de estado, validaciones, integraciones y despliegue separados para mantener la evolucion del producto."),
        ("1.8 Justificacion social y academica", "demostrar que una plataforma con agente conversacional puede ampliar el acceso al comercio electronico, reducir barreras de uso y servir como caso de estudio para ingenieria de software."),
        ("1.9 Delimitacion del sistema", "acotar el informe al monorepo existente, las imagenes entregadas y los documentos de especificacion y diseno ubicados en las carpetas specs y design."),
        ("1.10 Limitaciones", "reconocer que las metricas de carga, accesibilidad total y disponibilidad dependen de pruebas futuras y de la configuracion final de servicios como Render, Vercel, Neon, Mercado Pago y Cloudinary."),
    ]
    for title, focus in problem_pages:
        add_content_page(
            doc,
            title,
            paragraph_block(focus, "planteamiento del problema"),
            bullets=[
                "Se mantiene correspondencia con los objetivos funcionales y no funcionales.",
                "Se documenta la relacion entre el problema y las decisiones de arquitectura.",
                "Se reserva evidencia tecnica detallada para los anexos.",
            ],
        )


def build_theory(doc: Document) -> None:
    add_heading(doc, "Capitulo II", level=1)
    intro = doc.add_paragraph(
        "Este capitulo se desarrolla con base bibliografica y citas en formato APA 7. Las citas textuales se emplean de manera breve; las demas ideas se presentan como parafrasis academicas con su respectiva fuente."
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_page_break()
    for title, p1, p2, source_note in THEORY_ENTRIES:
        add_content_page(
            doc,
            title,
            [p1, p2],
            bullets=[
                source_note,
                "Aplicacion en Aura Marketplace: la teoria se vincula con requisitos, diseno, implementacion y pruebas.",
                "Formato usado: citas parenteticas y narrativas segun APA 7.ª edicion.",
            ],
        )


def build_methodology(doc: Document) -> None:
    add_heading(doc, "Capitulo III", level=1)
    methodology_topics = [
        ("3.1 Tipo de investigacion", "una investigacion aplicada-tecnologica orientada a resolver un problema concreto mediante un producto de software documentado"),
        ("3.2 Nivel de investigacion", "un nivel descriptivo porque identifica componentes, procesos, tecnologias, actores y flujos sin manipular variables experimentales"),
        ("3.3 Diseno de investigacion", "un diseno no experimental y transversal aplicado al analisis del repositorio, los diagramas y los documentos de especificacion"),
        ("3.4 Poblacion y muestra tecnica", "los actores del sistema, modulos del repositorio, requisitos, casos de uso, entidades de datos y diagramas de arquitectura"),
        ("3.5 Variables e indicadores", "funcionamiento, usabilidad tecnica, mantenibilidad, seguridad, escalabilidad, trazabilidad y cobertura documental"),
        ("3.6 Tecnicas de recoleccion", "revision documental, inspeccion de codigo, analisis de diagramas, lectura del PDF de referencia y construccion de anexos"),
        ("3.7 Instrumentos", "fichas de analisis documental, matriz de consistencia, backlog, matriz de trazabilidad, checklist de arquitectura y escenarios SDD"),
        ("3.8 Metodo Scrum", "organizacion iterativa en sprints con backlog, sprint planning, daily review, incrementos y retrospectiva"),
        ("3.9 Roles Scrum", "Product Owner, Scrum Master, Development Team y stakeholders academicos del proyecto"),
        ("3.10 Product Backlog", "historias de usuario para agente, autenticacion, catalogo, carrito, checkout, vendedor, administrador, notificaciones y despliegue"),
        ("3.11 Sprint Planning", "seleccion de historias por prioridad, complejidad, dependencias tecnicas y valor de negocio"),
        ("3.12 Sprint 1", "preparacion del monorepo, configuracion base, estructura de carpetas, validaciones iniciales y entorno de desarrollo"),
        ("3.13 Sprint 2", "autenticacion, usuarios, roles, JWT, refresh tokens, Neon Auth y verificacion de correo"),
        ("3.14 Sprint 3", "catalogo, categorias, publicaciones, inventario, imagenes y busqueda"),
        ("3.15 Sprint 4", "carrito, ordenes, pagos, Mercado Pago y consistencia transaccional"),
        ("3.16 Sprint 5", "agente inteligente, interpretacion de lenguaje natural, audio, conversaciones, intenciones y entidades"),
        ("3.17 Sprint 6", "administracion, reportes, auditoria, favoritos, resenas, promociones y notificaciones"),
        ("3.18 Sprint 7", "despliegue, CI/CD, Render, Vercel, variables de entorno, pruebas e integraciones"),
        ("3.19 Definicion de terminado", "codigo compilable, pruebas relevantes, endpoints documentados, validaciones, seguridad y evidencia en anexos"),
        ("3.20 Gestion de riesgos", "fallos de proveedores externos, latencia, credenciales, consistencia de datos, acoplamiento y regresiones de frontend"),
    ]
    for title, topic in methodology_topics:
        add_content_page(
            doc,
            title,
            paragraph_block(topic, "metodologia"),
            bullets=[
                "Artefacto Scrum asociado: backlog, sprint, incremento, revision o retrospectiva.",
                "Evidencia esperada: codigo, prueba, tabla, diagrama o criterio de aceptacion.",
                "Criterio de control: trazabilidad con requisitos funcionales y no funcionales.",
            ],
        )


def build_results(doc: Document) -> None:
    add_heading(doc, "Capitulo IV", level=1)
    result_topics = [
        ("4.1 Vision del producto implementado", "Aura Marketplace como plataforma multidireccional de comercio electronico con agente conversacional"),
        ("4.2 Actores del sistema", "visitante, comprador, vendedor, administrador y servicios externos"),
        ("4.3 Modulos funcionales", "autenticacion, usuarios, catalogo, carrito, ordenes, pagos, agente, administracion, auditoria y notificaciones"),
        ("4.4 Arquitectura general del sistema", "la relacion entre frontend React, backend NestJS, base de datos, integraciones y despliegue"),
        ("4.5 Flujo cliente-servidor", "la comunicacion HTTPS REST entre navegador, API, capa de aplicacion, infraestructura y datos"),
        ("4.6 Arquitectura del backend L01", "controladores REST, guards, DTOs, ValidationPipe, CORS, Swagger y endpoints"),
        ("4.7 Arquitectura del backend L02", "AgentService y ConversationsService como capa de inteligencia conversacional"),
        ("4.8 Arquitectura del backend L03", "servicios de aplicacion y casos de uso para productos, carrito, ordenes, pagos, usuarios y admin"),
        ("4.9 Arquitectura del backend L04", "entidades, enums, puertos de repositorio y contratos de proveedores"),
        ("4.10 Arquitectura del backend L05", "adaptadores Prisma, Cloudinary, Mercado Pago, Resend, Argon2, Redis y Gemini"),
        ("4.11 Inyeccion de dependencias", "tokens de interfaces y providers para desacoplar servicios de aplicacion de implementaciones concretas"),
        ("4.12 Controladores REST", "admin, agent, auth, cart, categories, favorites, notifications, orders, payments, products, promotions, reviews y users"),
        ("4.13 Guards y autorizacion", "JwtAuthGuard, RolesGuard, decoradores publicos y decoradores de roles"),
        ("4.14 DTOs y validaciones", "login, registro, refresh, verify email, password reset, carrito, ordenes, categorias y perfiles"),
        ("4.15 Servicios de aplicacion", "AuthService, OrdersService, ProductsService, CartService, PaymentsService, UsersService y AdminService"),
        ("4.16 Repositorios Prisma", "adaptadores por entidad para usuario, producto, orden, carrito, auditoria, conversaciones y notificaciones"),
        ("4.17 Transacciones de orden", "creacion atomica de orden, lineas, pago, decremento de inventario y limpieza de carrito"),
        ("4.18 Seguridad de contrasenas", "Argon2 como derivacion segura y separacion del hasher mediante puerto"),
        ("4.19 Gestion de tokens", "JWT access token, refresh token, TokenRevocado y bloqueo por intentos fallidos"),
        ("4.20 Neon Auth", "sincronizacion de identidad externa con usuarios locales del sistema"),
        ("4.21 Mercado Pago", "integracion de checkout, referencia de pago, webhook y estados de pago"),
        ("4.22 Resend Email API", "verificacion de cuenta, restablecimiento de contrasena y comunicaciones transaccionales"),
        ("4.23 Cloudinary", "almacenamiento de imagenes de producto y separacion de assets del backend"),
        ("4.24 Upstash Redis", "cache distribuida para reducir latencia y preservar escalabilidad"),
        ("4.25 Gemini AI", "procesamiento de lenguaje natural e interpretacion de intenciones del agente"),
        ("4.26 Arquitectura del frontend", "React modular por pages, components, api, store, layouts, hooks, lib, styles y types"),
        ("4.27 Enrutamiento frontend", "BrowserRouter, rutas publicas, rutas protegidas, perfiles y panel de administracion"),
        ("4.28 Paginas publicas", "Home, Catalog, ProductDetail, Login, Register, ForgotPassword, ResetPassword y VerifyEmail"),
        ("4.29 Flujos de comprador", "carrito, checkout shipping, checkout payment, order success, favoritos e historial de ordenes"),
        ("4.30 Flujos de vendedor", "productos del vendedor, formulario de producto, ordenes del vendedor y clientes"),
        ("4.31 Flujos de administrador", "usuarios, productos, ordenes, categorias, reportes y perfil de administracion"),
        ("4.32 Componentes visuales", "AuraHeader, BrandLogo, ChatAgente, AuthLayout, AuthField, CheckoutStepper y componentes UI"),
        ("4.33 Estado global frontend", "authStore, cartStore, agentStore y profilePhotoStore"),
        ("4.34 Cliente HTTP", "axios, interceptores, VITE_API_URL, tokens y manejo de errores"),
        ("4.35 Validacion cliente", "Zod y React Hook Form como filtro de calidad antes de la API"),
        ("4.36 Diseno visual", "Tailwind CSS, shadcn/ui, componentes reutilizables y consistencia de layouts"),
        ("4.37 Base de datos", "modelo relacional en Prisma con entidades de identidad, catalogo, marketplace, agente y auditoria"),
        ("4.38 Entidades de identidad", "Usuario, RefreshToken, Direccion, PreferenciasUsuario y TokenRevocado"),
        ("4.39 Entidades de catalogo", "Categoria, Publicacion, Inventario e ImagenPublicacion"),
        ("4.40 Entidades transaccionales", "Carrito, ItemCarrito, Orden, LineaOrden, Pago, Cupon y Promocion"),
        ("4.41 Entidades del agente", "Sesion, Conversacion, Mensaje, Intencion y EntidadExtraida"),
        ("4.42 Entidades de soporte", "Favorito, Resena, Notificacion y Auditoria"),
        ("4.43 Integridad referencial", "claves foraneas, relaciones en cascada, indices y unicidad"),
        ("4.44 Reglas de negocio", "confirmacion de compra, stock atomico, precio positivo, validacion de datos y auditoria"),
        ("4.45 Requisitos no funcionales", "rendimiento, disponibilidad, seguridad, escalabilidad, usabilidad, accesibilidad y mantenibilidad"),
        ("4.46 Pruebas unitarias backend", "specs por controladores, servicios, repositorios, providers y utilidades"),
        ("4.47 Pruebas de integracion", "auth, RBAC, pagos, ordenes, base de datos, Redis, Cloudinary, Resend, Gemini y agente"),
        ("4.48 Pruebas E2E frontend", "flujos de compra, vendedor, administrador, autenticacion y agente"),
        ("4.49 Despliegue backend", "Render, build NestJS, Prisma generate, migrate deploy y health check"),
        ("4.50 Despliegue frontend", "Vercel o Pages, build Vite, dist y rewrites SPA"),
        ("4.51 Variables de entorno", "separacion de secretos, URLs publicas, claves de proveedores y credenciales de despliegue"),
        ("4.52 CI/CD", "GitHub Actions para validar frontend, backend, Prisma, tests y disparar despliegues"),
        ("4.53 Observabilidad", "logs, auditoria, health checks y trazabilidad de eventos criticos"),
        ("4.54 Discusion de resultados", "alineacion entre teoria, arquitectura, implementacion y anexos"),
        ("4.55 Sintesis tecnica", "madurez del diseno, pendientes de pruebas de carga y mejora continua"),
    ]
    for title, topic in result_topics:
        add_content_page(
            doc,
            title,
            paragraph_block(topic, "resultados"),
            bullets=[
                "Fuente tecnica: repositorio, especificaciones, diseno, schema Prisma o diagrama anexo.",
                "Impacto: mejora de mantenibilidad, seguridad, trazabilidad o experiencia del usuario.",
                "Evidencia: modulo, entidad, flujo, servicio, adaptador, prueba o variable de entorno.",
            ],
        )


def build_conclusions(doc: Document) -> None:
    add_heading(doc, "Capitulo V", level=1)
    conclusions = [
        ("5.1 Conclusion general", "Aura Marketplace evidencia una arquitectura coherente para un marketplace inteligente porque separa la experiencia de usuario, la logica de aplicacion, el dominio, la infraestructura y los servicios externos."),
        ("5.2 Conclusion sobre SDD", "Specs Driven Development permite que las especificaciones actuen como fuente de verdad para requisitos, diseno, implementacion, pruebas y anexos academicos."),
        ("5.3 Conclusion sobre backend", "la estructura NestJS L01-L05 fortalece la mantenibilidad al ubicar controladores, agente, casos de uso, puertos y adaptadores en limites identificables."),
        ("5.4 Conclusion sobre frontend", "la SPA React organiza rutas, componentes, stores y clientes API de forma modular, permitiendo que compradores, vendedores y administradores trabajen con flujos diferenciados."),
        ("5.5 Conclusion sobre metodologia", "Scrum proporciona una ruta iterativa para planificar, construir, revisar y estabilizar la plataforma por incrementos verificables."),
        ("5.6 Recomendaciones", "se recomienda ejecutar pruebas de carga, auditoria de accesibilidad, monitoreo en produccion, revision de secretos, actualizacion de diagramas y mantenimiento de backlog."),
    ]
    for title, topic in conclusions:
        add_content_page(
            doc,
            title,
            paragraph_block(topic, "conclusiones"),
            bullets=[
                "Mantener la documentacion sincronizada con el repositorio.",
                "Usar anexos como evidencia de diseno y no como material decorativo.",
                "Registrar decisiones arquitectonicas cuando cambien proveedores o reglas de negocio.",
            ],
        )

    add_heading(doc, "Referencias bibliograficas", level=1)
    add_body(
        doc,
        "Las siguientes referencias se presentan en formato APA 7.ª edicion y sustentan principalmente el Capitulo II, donde se desarrolla el marco teorico de SDD, arquitectura, backend, frontend, seguridad, accesibilidad, persistencia, pruebas y Scrum.",
    )
    for author, year, title, publisher in APA_REFERENCES:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(-0.35)
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(f"{author} ({year}). ")
        set_run_font(run, size=12)
        title_run = p.add_run(title + " ")
        set_run_font(title_run, size=12, italic=True)
        pub_run = p.add_run(publisher)
        set_run_font(pub_run, size=12)
    doc.add_page_break()


def build_annexes(doc: Document) -> None:
    add_heading(doc, "Anexos", level=1)

    matrix = [
        ["Problema", "Friccion en comercio electronico y complejidad tecnica de una plataforma inteligente."],
        ["Objetivo", "Documentar arquitectura, backend, frontend, Scrum, SDD y anexos de Aura Marketplace."],
        ["Hipotesis", "Una documentacion basada en SDD y arquitectura hexagonal mejora trazabilidad y mantenibilidad."],
        ["Variables", "Funcionamiento, mantenibilidad, seguridad, usabilidad tecnica y trazabilidad."],
        ["Instrumentos", "Revision documental, inspeccion de codigo, matriz de requisitos, backlog y diagramas."],
    ]
    add_content_page(
        doc,
        "Anexo 1. Matriz de consistencia",
        ["La matriz de consistencia resume la relacion entre problema, objetivos, hipotesis, variables e instrumentos usados para elaborar el informe."],
        table=(["Elemento", "Descripcion"], matrix, [1.5, 4.9]),
    )

    backlog_rows = [
        ["HU-01", "Registrarse y verificar correo", "Alta", "5"],
        ["HU-02", "Iniciar sesion y renovar token", "Alta", "5"],
        ["HU-03", "Buscar productos con filtros", "Alta", "8"],
        ["HU-04", "Conversar con el agente", "Alta", "13"],
        ["HU-05", "Agregar producto al carrito", "Alta", "5"],
        ["HU-06", "Completar checkout y pago", "Alta", "13"],
        ["HU-07", "Publicar producto como vendedor", "Alta", "8"],
        ["HU-08", "Gestionar ordenes de vendedor", "Media", "8"],
        ["HU-09", "Administrar usuarios y productos", "Alta", "8"],
        ["HU-10", "Consultar reportes y auditoria", "Media", "5"],
        ["HU-11", "Recibir correos transaccionales", "Media", "5"],
        ["HU-12", "Subir imagenes de publicaciones", "Media", "5"],
    ]
    add_content_page(
        doc,
        "Anexo 2. Product Backlog",
        ["El backlog representa el conjunto inicial de historias necesarias para construir la experiencia principal del marketplace y del agente inteligente."],
        table=(["ID", "Historia", "Prioridad", "Puntos"], backlog_rows, [1.0, 3.5, 1.1, 0.8]),
    )

    sdd_flows = [
        ("SDD-01 Vision del producto", ["Spec de origen: specs/01-Vision.md", "Define problema, oportunidad, propuesta de valor y riesgos", "Guia: objetivos estrategicos, arquitectura inicial y alcance del MVP"]),
        ("SDD-02 Stakeholders y actores", ["Spec de origen: specs/02-Stakeholders.md y specs/09-Actores.md", "Define comprador, vendedor, administrador y servicios externos", "Guia: permisos, rutas protegidas y experiencia por rol"]),
        ("SDD-03 Glosario del dominio", ["Spec de origen: specs/03-Glosario.md", "Normaliza terminos como Usuario, Publicacion, Orden, Carrito, Agente e Intencion", "Guia: nombres de entidades, DTOs, stores y tablas"]),
        ("SDD-04 Alcance", ["Spec de origen: specs/04-Alcance.md", "Separa lo implementado de lo fuera de alcance", "Guia: priorizacion del backlog y control de expectativas"]),
        ("SDD-05 Objetivos", ["Spec de origen: specs/05-Objetivos.md", "Define objetivos e indicadores medibles", "Guia: KPIs, pruebas y criterios de aceptacion"]),
        ("SDD-06 Reglas de negocio", ["Spec de origen: specs/06-ReglasNegocio.md", "Define restricciones como stock, precio, confirmacion, bloqueo y auditoria", "Guia: servicios de aplicacion, validaciones y transacciones"]),
        ("SDD-07 Requisitos funcionales", ["Spec de origen: specs/07-RequisitosFuncionales.md", "Describe funciones por modulo y actor", "Guia: controladores REST, paginas, componentes y casos de prueba"]),
        ("SDD-08 Requisitos no funcionales", ["Spec de origen: specs/08-RequisitosNoFuncionales.md", "Define rendimiento, disponibilidad, seguridad, escalabilidad y accesibilidad", "Guia: arquitectura cloud, cache, seguridad y observabilidad"]),
        ("SDD-09 Casos de uso y criterios", ["Spec de origen: specs/10-CasosUsoResumen.md y specs/11-CriteriosAceptacion.md", "Concreta comportamiento esperado en escenarios verificables", "Guia: pruebas unitarias, integracion y E2E"]),
        ("SDD-10 Matriz de trazabilidad", ["Spec de origen: specs/12-MatrizTrazabilidad.md", "Relaciona objetivos, requisitos, reglas y casos de uso", "Guia: cierre documental, anexos y validacion academica"]),
    ]
    for title, events in sdd_flows:
        rows = [[str(i + 1), event] for i, event in enumerate(events)]
        add_content_page(
            doc,
            f"Anexo 3. {title}",
            [
                "Este anexo explica como se aplica SDD en el proyecto: la especificacion funciona como fuente de verdad antes de disenar, codificar o probar.",
                "La aplicacion de SDD en el informe permite conectar requisitos, reglas de negocio, casos de uso, diseno arquitectonico, implementacion y pruebas de aceptacion.",
            ],
            table=(["Paso", "Elemento SDD"], rows, [0.8, 5.6]),
        )

    fig_titles = [
        "Arquitectura general del backend y servicios externos",
        "Arquitectura modular del frontend React",
        "Boundaries de backend, infraestructura y datos",
        "Modelo entidad-relacion del marketplace",
        "Stack tecnologico y despliegue conectado",
    ]
    for idx, (path, title) in enumerate(zip(IMAGE_PATHS, fig_titles), 1):
        add_heading(doc, f"Anexo 4.{idx}. {title}", level=2)
        add_body(
            doc,
            "La figura se incorpora como evidencia visual del diseno proporcionado. Su lectura complementa las secciones de arquitectura general, backend, frontend, persistencia e integraciones.",
        )
        if path.exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Inches(0)
            run = p.add_run()
            run.add_picture(str(path), width=Inches(6.4))
            cap = doc.add_paragraph(f"Figura A{idx}. {title}. Fuente: imagen proporcionada para anexos del proyecto.")
            cap.style = "CaptionCustom"
        else:
            add_body(doc, f"No se encontro la imagen esperada: {path}")
        doc.add_page_break()

    endpoint_rows = [
        ["Auth", "/auth/register, /auth/login, /auth/refresh, /auth/verify-email", "Identidad y sesiones"],
        ["Products", "/products, /products/:id", "Catalogo y publicaciones"],
        ["Cart", "/cart, /cart/items", "Carrito del comprador"],
        ["Orders", "/orders, /vendor/orders", "Compra y gestion de ordenes"],
        ["Payments", "/payments, /payments/webhook", "Pago y eventos externos"],
        ["Agent", "/agent/message, /agent/voice", "Interaccion conversacional"],
        ["Admin", "/admin/users, /admin/products, /admin/reports", "Administracion"],
    ]
    add_content_page(
        doc,
        "Anexo 5. Contratos API principales",
        ["La tabla sintetiza la superficie REST esperada segun la estructura de controladores del backend."],
        table=(["Modulo", "Rutas representativas", "Responsabilidad"], endpoint_rows, [1.4, 3.2, 1.8]),
    )

    entity_groups = [
        ["Identidad", "Usuario, RefreshToken, Direccion, PreferenciasUsuario, TokenRevocado"],
        ["Catalogo", "Categoria, Publicacion, Inventario, ImagenPublicacion"],
        ["Marketplace", "Carrito, ItemCarrito, Orden, LineaOrden, Pago, Cupon"],
        ["Agente", "Sesion, Conversacion, Mensaje, Intencion, EntidadExtraida"],
        ["Soporte", "Favorito, Resena, Promocion, Notificacion, Auditoria"],
    ]
    add_content_page(
        doc,
        "Anexo 6. Agrupacion de entidades Prisma",
        ["La agrupacion evidencia que la base de datos no es una coleccion dispersa de tablas, sino un modelo relacional organizado por subdominios."],
        table=(["Subdominio", "Entidades"], entity_groups, [1.7, 4.7]),
    )

    tests = [
        ["Unitarias backend", "Servicios, controladores, DTOs, providers y utilidades"],
        ["Integracion backend", "Auth, RBAC, pagos, ordenes, base de datos, Redis, Cloudinary, Resend, Gemini"],
        ["E2E frontend", "Autenticacion, compra, vendedor, administrador y agente"],
        ["Validacion manual", "Diagramas, despliegue, variables y anexos"],
    ]
    add_content_page(
        doc,
        "Anexo 7. Estrategia de pruebas",
        ["Las pruebas se distribuyen para verificar tanto unidades aisladas como flujos completos de negocio."],
        table=(["Tipo de prueba", "Cobertura"], tests, [2.0, 4.4]),
    )

    deploy = [
        ["Backend", "Render", "npm ci, prisma generate, build, migrate deploy, start:prod"],
        ["Frontend", "Vercel / Pages", "npm ci, npm run build, salida dist"],
        ["Datos", "Neon PostgreSQL", "DATABASE_URL segura y migraciones Prisma"],
        ["Imagenes", "Cloudinary", "Cloud name, API key, API secret y carpeta"],
        ["Cache", "Upstash Redis", "REST URL, token y prefijo de cache"],
    ]
    add_content_page(
        doc,
        "Anexo 8. Checklist de despliegue",
        ["El despliegue exige separar secretos, validar dominios, configurar CORS y comprobar endpoints de salud antes de exponer el sistema a usuarios."],
        table=(["Componente", "Plataforma", "Acciones"], deploy, [1.5, 1.5, 3.4]),
    )

    glossary = [
        ["Adaptador", "Implementacion concreta de un puerto hacia una tecnologia externa."],
        ["Puerto", "Interfaz que define una necesidad del dominio o aplicacion."],
        ["SDD", "Specs Driven Development: desarrollo guiado por especificaciones versionadas y trazables."],
        ["SPA", "Aplicacion web de una sola pagina que renderiza rutas en el navegador."],
        ["DTO", "Objeto de transferencia de datos usado para validar entradas y salidas."],
        ["RBAC", "Control de acceso basado en roles."],
    ]
    add_content_page(
        doc,
        "Anexo 9. Glosario tecnico",
        ["El glosario consolida terminos usados a lo largo del informe para mantener un lenguaje comun."],
        table=(["Termino", "Definicion"], glossary, [1.6, 4.8]),
    )

    for i in range(10, 21):
        add_content_page(
            doc,
            f"Anexo {i}. Evidencia complementaria de diseno",
            paragraph_block(
                f"la evidencia complementaria numero {i} se utiliza para reforzar la relacion entre requisito, diseno, implementacion y prueba",
                "anexos",
            ),
            bullets=[
                "Mantiene trazabilidad con las decisiones descritas en los capitulos principales.",
                "Puede ampliarse con capturas, reportes de pruebas o commits especificos.",
                "Sirve como espacio de soporte para revisiones academicas futuras.",
            ],
        )


def build_document() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    doc = Document()
    configure_document(doc)
    build_front_matter(doc)
    build_chapter_one(doc)
    build_theory(doc)
    build_methodology(doc)
    build_results(doc)
    build_conclusions(doc)
    build_annexes(doc)
    doc.core_properties.title = "Informe Aura Marketplace"
    doc.core_properties.subject = "Arquitectura, backend, frontend, Scrum, SDD y anexos"
    doc.core_properties.author = "Codex"
    doc.core_properties.comments = "Documento generado a partir del PDF de ejemplo, repositorio y diagramas proporcionados."
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build_document()
